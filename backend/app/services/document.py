import uuid
import mimetypes
from io import BytesIO

from sqlalchemy.orm import Session
from sqlalchemy import func, text
from openai import OpenAI
from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl

from app.models.document import Document, DocumentChunk
from app.schemas.documents.requests import UploadDocumentRequest, FAQSearchRequest
from app.services.s3 import s3_service
from app.utils.responses import ResponseHandler
from app.core.config import settings

openai_client = OpenAI(api_key=settings.OPENAI_API_KEY)

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50


class DocumentService:

    @staticmethod
    def upload_document(
        db: Session,
        file_bytes: bytes,
        filename: str,
        payload: UploadDocumentRequest
    ):
        """Upload a new document or a new version of an existing document group."""
        file_type = filename.rsplit(".", 1)[-1].lower()
        content_type = mimetypes.guess_type(
            filename)[0] or "application/octet-stream"

        # 1. Upload to S3
        s3_key = f"documents/{uuid.uuid4()}/{filename}"
        s3_url = s3_service.upload_file(
            filename=s3_key,
            file_content=file_bytes,
            content_type=content_type
        )

        # 2. Resolve group & version
        if payload.document_group_id:
            DocumentService._deactivate_group(db, payload.document_group_id)
            latest_version = db.query(func.max(Document.version)).filter(
                Document.document_group_id == payload.document_group_id
            ).scalar() or 0
            document_group_id = payload.document_group_id
        else:
            document_group_id = uuid.uuid4()
            latest_version = 0

        # 3. Create document record
        doc = Document(
            title=payload.title,
            description=payload.description,
            file_type=file_type,
            s3_key=s3_key,
            s3_url=s3_url,
            document_group_id=document_group_id,
            version=latest_version + 1,
            is_active=True
        )
        db.add(doc)
        db.flush()

        # 4. Parse -> Chunk -> Embed -> Save
        raw_text = DocumentService._parse_file(file_bytes, file_type)
        for i, chunk_content in enumerate(DocumentService._chunk_text(raw_text)):
            db.add(DocumentChunk(
                document_id=doc.id,
                content=chunk_content,
                embedding=DocumentService._embed(chunk_content),
                chunk_index=i,
                is_active=True
            ))

        db.commit()
        db.refresh(doc)

        return ResponseHandler.success(
            message="Document uploaded and indexed successfully",
            data=doc
        )

    @staticmethod
    def get_all_documents(db: Session):
        """List all active documents."""
        docs = db.query(Document).filter(
            Document.is_active == True
        ).order_by(Document.created_at.desc()).all()

        return ResponseHandler.success(
            message="Documents retrieved successfully",
            data=docs
        )

    @staticmethod
    def delete_document(db: Session, document_id: uuid.UUID):
        """Deactivate a document and all its chunks."""
        doc = db.query(Document).filter(
            Document.id == document_id,
            Document.is_active == True
        ).first()

        if not doc:
            ResponseHandler.not_found_error("Document", str(document_id))

        doc.is_active = False
        db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).update({"is_active": False})

        db.commit()

        return ResponseHandler.success(
            message="Document deactivated successfully",
            data=None
        )

    @staticmethod
    def search_faq(db: Session, payload: FAQSearchRequest):
        """Semantic search over active document chunks."""
        embedding_str = DocumentService._embedding_to_str(
            DocumentService._embed(payload.query)
        )

        results = db.execute(
            text("""
                SELECT
                    dc.content,
                    d.title                                     AS document_title,
                    d.version,
                    1 - (dc.embedding <=> :embedding ::vector)  AS similarity
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
                WHERE dc.is_active = true
                  AND d.is_active  = true
                ORDER BY dc.embedding <=> :embedding ::vector
                LIMIT :limit
            """),
            {"embedding": embedding_str, "limit": payload.limit}
        ).fetchall()

        return ResponseHandler.success(
            message="FAQ search completed",
            data=[dict(row._mapping) for row in results]
        )

    #  Private helpers

    @staticmethod
    def _deactivate_group(db: Session, document_group_id: uuid.UUID):
        """Deactivate all documents and chunks in a group."""
        doc_ids = [
            row.id for row in db.query(Document.id).filter(
                Document.document_group_id == document_group_id
            ).all()
        ]

        if doc_ids:
            db.query(DocumentChunk).filter(
                DocumentChunk.document_id.in_(doc_ids)
            ).update({"is_active": False}, synchronize_session="fetch")

        db.query(Document).filter(
            Document.document_group_id == document_group_id
        ).update({"is_active": False}, synchronize_session="fetch")

    @staticmethod
    def _parse_file(file_bytes: bytes, file_type: str) -> str:
        """Extract plain text from PDF, DOCX, or XLSX."""
        if file_type == "pdf":
            reader = PdfReader(BytesIO(file_bytes))
            return "\n".join(
                page.extract_text() for page in reader.pages if page.extract_text()
            )

        if file_type == "docx":
            doc = DocxDocument(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        if file_type in ("xlsx", "xls"):
            wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
            rows = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    line = " | ".join(str(c) for c in row if c is not None)
                    if line:
                        rows.append(line)
            return "\n".join(rows)

        ResponseHandler.bad_request(f"Unsupported file type: {file_type}")

    @staticmethod
    def _chunk_text(raw_text: str) -> list[str]:
        """Split text into overlapping chunks."""
        chunks, start = [], 0
        while start < len(raw_text):
            chunk = raw_text[start: start + CHUNK_SIZE].strip()
            if chunk:
                chunks.append(chunk)
            start += CHUNK_SIZE - CHUNK_OVERLAP
        return chunks

    @staticmethod
    def _embed(text: str) -> list[float]:
        """Call OpenAI embeddings API."""
        response = openai_client.embeddings.create(
            input=text,
            model="text-embedding-ada-002"
        )
        return response.data[0].embedding

    @staticmethod
    def _embedding_to_str(embedding: list[float]) -> str:
        """Convert embedding list to pgvector string format."""
        return f"[{','.join(str(x) for x in embedding)}]"
